"""Generate Word documentation for the Healthy Meal Recommendation System."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            tbl.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()


# ═══════════════  TITLE  ═══════════════
title = doc.add_heading("Healthy Meal Recommendation System", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Project Documentation")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# ═══════════════  1. OVERVIEW  ═══════════════
doc.add_heading("1. Project Overview", level=1)
add_table(
    ["Item", "Details"],
    [
        ["Project Name", "Healthy Meal Recommendation System"],
        ["Type", "Graduation Project"],
        ["Stack", "Python - Scikit-learn - Streamlit - FastAPI"],
        ["AI Model", "Multi-Output Regression (Random Forest / XGBoost / LightGBM)"],
        ["Chatbot", "TF-IDF search + OpenRouter GPT-4o-mini fallback"],
        ["Data", "501-row nutrition CSV + 64-recipe JSON"],
    ],
)

# ═══════════════  2. ARCHITECTURE  ═══════════════
doc.add_heading("2. System Architecture", level=1)
doc.add_paragraph(
    "The system consists of four main layers:\n"
    "  - Streamlit UI (streamlit_app.py) -- user-facing web interface with two tabs\n"
    "  - FastAPI (api.py) -- REST API for external integrations\n"
    "  - ML Engine (meal_model.py) -- trains and serves nutrition predictions + recipe recommendations\n"
    "  - Chatbot Engine (chatbot_engine.py) -- TF-IDF recipe search with LLM-enhanced answers"
)
doc.add_paragraph(
    "Data Flow:\n"
    "  1. User fills profile -> ML model predicts daily nutrition targets\n"
    "  2. RecipeRecommender filters 64 recipes by allergies, diet, health conditions\n"
    "  3. Remaining recipes are scored against per-meal nutrition targets\n"
    "  4. Top 5 meals returned with scores, reasons, and nutritional info"
)

# ═══════════════  3. FILE BREAKDOWN  ═══════════════
doc.add_heading("3. File-by-File Breakdown", level=1)

doc.add_heading("3.1 meal_model.py -- Core ML + Recommendation Engine (433 lines)", level=2)
add_table(
    ["Component", "Description"],
    [
        ["UserProfile", "Dataclass: age, gender, weight, height, allergies, conditions, diet, goals, nutrition limits"],
        ["UserNutritionModel", "Scikit-learn Pipeline (preprocessor + regressor) to predict daily calorie, protein, carbs, and fat targets"],
        ["RecipeRecommender", "Filters and scores 64 recipes against per-meal nutrition targets, respecting allergies, diet, health conditions"],
        ["MealRecommendationSystem", "Facade: loads model + recommender, runs predict -> recommend pipeline"],
    ],
)
doc.add_paragraph("ML Pipeline Steps:")
doc.add_paragraph("  1. Numeric features (Age, Height, Weight) -> median impute")
doc.add_paragraph("  2. Categorical features (Gender, Activity Level, Fitness Goal, Dietary Preference) -> OneHotEncoder")
doc.add_paragraph("  3. Model: MultiOutputRegressor wrapping Random Forest, XGBoost, or LightGBM")
doc.add_paragraph()
p = doc.add_paragraph("Recipe Scoring: score = 0.35*cal + 0.30*protein + 0.20*carbs + 0.15*fats + bonus")
p.runs[0].italic = True

doc.add_heading("3.2 chatbot_engine.py -- Chatbot Engine (316 lines)", level=2)
add_table(
    ["Component", "Description"],
    [
        ["FoodChatbot", "TF-IDF-based recipe search + intent classification (greeting/thanks/help)"],
        ["Search", "TF-IDF index over recipe names, ingredients, categories, diet tags, diseases"],
        ["Filtering", "Respects user allergies, diet type, and health conditions"],
        ["LLM", "OpenRouter API (GPT-4o-mini) with context; falls back to local rule-based answers"],
    ],
)

doc.add_heading("3.3 streamlit_app.py -- Web UI (358 lines)", level=2)
add_table(
    ["Tab", "Features"],
    [
        ["Recommendations", "User profile form -> ML prediction -> Top 5 meal cards with scores, macros, images, ingredients"],
        ["Chatbot", "Chat interface -> searches recips.json -> answers via OpenRouter or local fallback"],
    ],
)

doc.add_heading("3.4 api.py -- REST API (217 lines)", level=2)
add_table(
    ["Endpoint", "Method", "Description"],
    [
        ["/", "GET", "API status"],
        ["/health", "GET", "Health check"],
        ["/predict-targets", "POST", "Predict daily nutrition targets"],
        ["/recommend", "POST", "Full recommendation pipeline"],
        ["/chat", "POST", "Chatbot conversation"],
        ["/recipes", "GET", "List recipes"],
    ],
)

doc.add_heading("3.5 train_model.py -- Training Script (9 lines)", level=2)
doc.add_paragraph("Trains UserNutritionModel on nutrition_dataset.csv and saves to nutrition_model.joblib.")

doc.add_heading("3.6 Data Files", level=2)
add_table(
    ["File", "Size", "Content"],
    [
        ["nutrition_dataset.csv", "102 KB", "501 rows of user profiles with nutrition targets"],
        ["recips.json", "54 KB", "64 structured recipes with macros, diet, allergy, disease info"],
        ["nutrition_model.joblib", "~15 MB", "Pre-trained ML model pipeline"],
    ],
)

# ═══════════════  4. DATA QUALITY  ═══════════════
doc.add_heading("4. Data Quality Issues", level=1)
p = doc.add_paragraph("WARNING: Dataset has structural problems that should be fixed:")
p.runs[0].bold = True
p.runs[0].font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
add_table(
    ["Issue", "Details"],
    [
        ["Repeated headers", "Lines 42 and 210 contain header rows mixed into the data"],
        ["Inconsistent columns", "\"Protein\" vs \"Protein (g)\", \"Fat\" vs \"Fat (g)\" in different sections"],
        ["Typo in file name", "File named recips.json instead of recipes.json"],
        ["Typo in allergy", "Some recipes use \"lactuse\" instead of \"lactose\""],
        ["Small dataset", "Only 501 rows, many near-duplicates"],
        ["Small recipe pool", "Only 64 recipes limits recommendation variety"],
    ],
)

# ═══════════════  5. WHAT TO ADD  ═══════════════
doc.add_heading("5. What to ADD (Recommendations)", level=1)

items_to_add = [
    ("5.1 Model Evaluation & Comparison Script",
     "Why: No MAE/RMSE/R2 metrics are logged. No way to evaluate model quality.",
     "Action: Add cross-validation + metrics logging. Save results to model_comparison_results.json."),
    ("5.2 User Authentication & Profile Saving",
     "Why: Users must re-enter their profile every time.",
     "Action: Use streamlit-authenticator or a simple SQLite database."),
    ("5.3 Meal Planning (Daily / Weekly)",
     "Why: Currently only shows individual meals. A full day plan is more useful.",
     "Action: Group recipes by meal_type. Ensure total daily macros stay within targets."),
    ("5.4 More Recipes",
     "Why: 64 recipes is very limited.",
     "Action: Expand to 200+ recipes or integrate an external recipe API."),
    ("5.5 Arabic Language Support",
     "Why: UI is fully English but target users may be Arabic speakers.",
     "Action: Translate Streamlit labels, buttons, and headers."),
    ("5.6 Unit Tests",
     "Why: Zero tests exist. Tests demonstrate code quality.",
     "Action: Test model predictions, recipe filtering, and chatbot intents."),
    ("5.7 Logging & Error Handling",
     "Why: All except Exception: pass blocks silently swallow errors.",
     "Action: Add Python logging module."),
    ("5.8 .env Support for API Keys",
     "Why: st.secrets only works in Streamlit. FastAPI has no key management.",
     "Action: Add python-dotenv for unified config."),
    ("5.9 Docker Deployment",
     "Why: Makes the project easily deployable.",
     "Action: Add Dockerfile + docker-compose.yml."),
    ("5.10 Model Explainability",
     "Why: Users don't know why a calorie target was predicted.",
     "Action: Add SHAP values or feature importance visualization."),
]

for t, why, action in items_to_add:
    doc.add_heading(t, level=2)
    doc.add_paragraph(why)
    p_a = doc.add_paragraph(action)
    p_a.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ═══════════════  6. WHAT TO REMOVE  ═══════════════
doc.add_heading("6. What to REMOVE or REFACTOR", level=1)

items_to_remove = [
    ("6.1 Duplicated Code Between Files",
     "Problem: normalize_text, allergy checking, diet matching duplicated across streamlit_app.py, chatbot_engine.py, meal_model.py.",
     "Action: Consolidate into a single utils.py module (~150 lines saved)."),
    ("6.2 Duplicated OpenRouter API Call",
     "Problem: LLM call implemented twice in streamlit_app.py and chatbot_engine.py.",
     "Action: Use only FoodChatbot from chatbot_engine.py in Streamlit."),
    ("6.3 __pycache__ Directory",
     "Problem: Compiled cache files in the project. Should be in .gitignore.",
     "Action: Add to .gitignore, delete from repository."),
    ("6.4 Header Rows in Dataset",
     "Problem: Lines 42 and 210 are repeated header rows.",
     "Action: Clean the CSV once. Remove runtime filtering code."),
    ("6.5 Unused best_nutrition_model.joblib Path",
     "Problem: streamlit_app.py references a file that doesn't exist.",
     "Action: Produce it via model comparison or remove the reference."),
    ("6.6 Overly Broad Exception Handling",
     "Problem: except Exception: pass blocks swallow all errors.",
     "Action: Replace pass with logging.exception(...)."),
    ("6.7 Hardcoded API Model",
     "Problem: \"openai/gpt-4o-mini\" hardcoded in two places.",
     "Action: Move to config variable or environment variable."),
]

for t, problem, action in items_to_remove:
    doc.add_heading(t, level=2)
    doc.add_paragraph(problem)
    p_a = doc.add_paragraph(action)
    p_a.runs[0].font.color.rgb = RGBColor(0xCC, 0x33, 0x00)

# ═══════════════  7. SUMMARY TABLE  ═══════════════
doc.add_heading("7. Priority Summary", level=1)
add_table(
    ["Category", "Item", "Priority"],
    [
        ["Add", "Model evaluation metrics", "High"],
        ["Add", "Unit tests", "High"],
        ["Add", "More recipes (200+)", "High"],
        ["Add", "Daily meal planning", "Medium"],
        ["Add", "Logging & error handling", "Medium"],
        ["Add", "Arabic UI support", "Medium"],
        ["Add", "User profiles / auth", "Low"],
        ["Add", "Docker deployment", "Low"],
        ["Add", "Model explainability (SHAP)", "Low"],
        ["Remove", "Duplicated code across 3 files", "High"],
        ["Remove", "Duplicated OpenRouter call", "Medium"],
        ["Remove", "__pycache__ from repo", "Medium"],
        ["Fix", "Header rows in CSV dataset", "Medium"],
        ["Fix", "except Exception: pass blocks", "Medium"],
        ["Fix", "Typo: recips.json -> recipes.json", "Low"],
        ["Fix", "Typo: lactuse -> lactose in data", "Low"],
    ],
)

import os
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Project_Documentation.docx")
doc.save(OUTPUT)
print(f"Document saved to: {OUTPUT}")
